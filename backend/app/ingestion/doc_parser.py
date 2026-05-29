"""Parse uploaded strategic-goals documents and KB pillar docs (PDF / Word / text)."""
from __future__ import annotations

import io
from pathlib import Path

from app.analysis.schemas import GoalsDocument


def _parse_pdf(data: bytes) -> tuple[str, list[str]]:
    import fitz  # PyMuPDF

    sections: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text").strip()
            if text:
                sections.append(text)
    return "\n\n".join(sections), sections


def _parse_docx(data: bytes) -> tuple[str, list[str]]:
    from docx import Document

    document = Document(io.BytesIO(data))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Treat heading-styled paragraphs as section boundaries.
    sections: list[str] = []
    current: list[str] = []
    for p in document.paragraphs:
        if p.style and p.style.name and p.style.name.lower().startswith("heading"):
            if current:
                sections.append("\n".join(current))
                current = []
        if p.text.strip():
            current.append(p.text.strip())
    if current:
        sections.append("\n".join(current))
    return "\n".join(paragraphs), sections or paragraphs


def parse_document(filename: str, data: bytes) -> GoalsDocument:
    """Parse raw file bytes into normalized text + coarse sections."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text, sections = _parse_pdf(data)
    elif suffix in (".docx", ".doc"):
        text, sections = _parse_docx(data)
    elif suffix in (".txt", ".md"):
        text = data.decode("utf-8", errors="replace")
        sections = [b for b in text.split("\n\n") if b.strip()]
    else:
        raise ValueError(f"Unsupported document type: {suffix or '(none)'}")
    return GoalsDocument(filename=filename, text=text.strip(), sections=sections)


def parse_path(path: Path) -> GoalsDocument:
    return parse_document(path.name, path.read_bytes())
